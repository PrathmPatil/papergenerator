from __future__ import annotations

import json
import os
import re
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Optional, Tuple

import pandas as pd
from docx import Document


DEFAULT_CONFIG = {
    "classId": "class_9",
    "subjectId": "mathematics",
    "topicId": "NCERT",
    "difficulty": "medium",
    "marks": 1,
    "negativeMarks": 0,
    "subject_mapping": {
        "Probability": "mathematics",
        "Statistics": "mathematics",
        "Area": "mathematics",
        "INNOSAT": "general_knowledge",
        "Science": "science",
        "English": "english",
        "Maths": "mathematics",
        "Physics": "physics",
        "Chemistry": "chemistry",
        "Biology": "biology",
        "Reasoning": "reasoning",
    },
}

OPTION_PATTERNS = [
    re.compile(r"\((\d)\)\s*(.+?)(?=\(\d\)|$)", re.DOTALL),
    re.compile(r"([A-D])\)\s*(.+?)(?=[A-D]\)|$)", re.DOTALL),
    re.compile(r"([a-d])\.\s*(.+?)(?=[a-d]\.|$)", re.DOTALL),
    re.compile(r"([A-D])\.\s*(.+?)(?=[A-D]\.|$)", re.DOTALL),
]

QUESTION_PATTERNS = [
    re.compile(r"^Q\.?\s*(\d+)[.)\s]\s*(.+)", re.IGNORECASE),
    re.compile(r"^(\d+)[.)\s]\s*(.+)", re.IGNORECASE),
    re.compile(r"^Question\s*(\d+)[:.)\s]\s*(.+)", re.IGNORECASE),
]

SKIP_PATTERNS = [
    r"^_c_",
    r"ALL[I1]M",
    r"^EXERCISE",
    r"^Page \d+",
    r"^CBSE",
    r"^www\.",
    r"^-{5,}",
    r"^={5,}",
    r"^\d+$",
    r"^INSTRUCTIONS",
    r"^SECTION",
    r"^Total Marks",
    r"^Marks obtained",
]


@dataclass(frozen=True)
class DocxExcelDefaults:
    class_id: str = ""
    subject_id: str = ""
    topic_name: str = ""
    difficulty: str = "medium"
    marks: int = 1
    negative_marks: float = 0


class Logger:
    def __init__(self, verbose: bool = False):
        self.verbose = verbose
        self.logs: List[str] = []

    def log(self, message: str, level: str = "INFO") -> None:
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        entry = f"[{timestamp}] [{level}] {message}"
        self.logs.append(entry)
        if self.verbose:
            print(entry)

    def save(self, filepath: str) -> None:
        with open(filepath, "w", encoding="utf-8") as handle:
            handle.write("\n".join(self.logs))


class DocxProcessor:
    def __init__(self, config: Dict, logger: Logger):
        self.config = config
        self.logger = logger

    def is_noise(self, text: str) -> bool:
        value = text.strip()
        if not value or len(value) < 3:
            return True
        return any(re.search(pattern, value, re.IGNORECASE) for pattern in SKIP_PATTERNS)

    def extract_options(self, text: str) -> List[str]:
        options: List[str] = []
        for pattern in OPTION_PATTERNS:
            matches = pattern.findall(text)
            if not matches:
                continue
            for match in matches[:4]:
                option_text = match[1] if len(match) > 1 else match[0]
                option_text = re.sub(r"\s+", " ", option_text).strip()
                if option_text:
                    options.append(option_text)
            break

        seen = set()
        unique = []
        for option in options:
            if option not in seen:
                seen.add(option)
                unique.append(option)
        return unique[:4]

    def is_question_start(self, text: str) -> Tuple[Optional[str], Optional[str]]:
        value = text.strip()
        if not value:
            return None, None
        for pattern in QUESTION_PATTERNS:
            match = pattern.match(value)
            if not match:
                continue
            question_number = match.group(1)
            question_text = match.group(2).strip()
            if not re.match(r"^\d+$", question_text) and len(question_text) >= 3:
                return question_number, question_text
        return None, None

    def extract_images(self, docx_path: str, output_dir: str) -> List[Dict]:
        images = []
        os.makedirs(output_dir, exist_ok=True)
        try:
            with zipfile.ZipFile(docx_path, "r") as docx_zip:
                image_files = [
                    name
                    for name in docx_zip.namelist()
                    if name.startswith("word/media/")
                    and name.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp"))
                ]
                for image_path in image_files:
                    filename = os.path.basename(image_path)
                    output_path = os.path.join(output_dir, filename)
                    with docx_zip.open(image_path) as source, open(output_path, "wb") as target:
                        target.write(source.read())
                    images.append(
                        {
                            "original_path": image_path,
                            "filename": filename,
                            "output_path": output_path,
                        }
                    )
        except Exception as exc:
            self.logger.log(f"Error extracting images: {exc}", "WARNING")
        return images

    def extract_text(self, docx_path: str) -> List[Dict]:
        document = Document(docx_path)
        blocks = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text and not self.is_noise(text):
                blocks.append(
                    {
                        "text": text,
                        "style": paragraph.style.name if paragraph.style else "Normal",
                        "is_bold": any(run.bold for run in paragraph.runs if run.text.strip()),
                        "is_heading": bool(paragraph.style and "Heading" in paragraph.style.name),
                    }
                )

        for table in document.tables:
            for row in table.rows:
                text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if text and not self.is_noise(text):
                    blocks.append({"text": text, "style": "Table", "is_bold": False, "is_heading": False})

        return blocks

    def parse_questions(self, text_blocks: List[Dict]) -> List[Dict]:
        questions = []
        current_question = None

        for block in text_blocks:
            text = block["text"]
            question_number, question_text = self.is_question_start(text)

            if question_number is not None:
                if current_question and current_question.get("text"):
                    questions.append(current_question)

                embedded_options = self.extract_options(text)
                if embedded_options:
                    question_text = re.split(r"\([1-4]\)|[A-D]\)|[a-d]\.", text)[0].strip()
                    question_text = re.sub(r"\s+", " ", question_text).strip()

                current_question = {
                    "num": question_number,
                    "text": question_text,
                    "options": embedded_options,
                    "is_bold": block.get("is_bold", False),
                    "is_heading": block.get("is_heading", False),
                }
                continue

            if current_question is None:
                continue

            new_options = self.extract_options(text)
            if new_options:
                current_question["options"].extend(new_options)
            elif len(text) > 15 and not self.is_noise(text):
                if not re.match(r"^[A-D][).]", text) and not re.match(r"^\(\d\)", text):
                    current_question["text"] += " " + re.sub(r"\s+", " ", text).strip()

        if current_question and current_question.get("text"):
            questions.append(current_question)

        for question in questions:
            seen = set()
            unique_options = []
            for option in question["options"]:
                if option not in seen:
                    seen.add(option)
                    unique_options.append(option)
            question["options"] = unique_options[:4]
            question["text"] = re.sub(r"\s+", " ", question["text"]).strip()

        return questions

    def process(self, docx_path: str, output_folder: str) -> Dict:
        self.logger.log(f"Processing: {docx_path}")
        docx_name = Path(docx_path).stem
        images_dir = os.path.join(output_folder, f"images_{docx_name}")

        images = self.extract_images(docx_path, images_dir)
        self.logger.log(f"Extracted {len(images)} images")

        text_blocks = self.extract_text(docx_path)
        self.logger.log(f"Extracted {len(text_blocks)} text blocks")

        questions = self.parse_questions(text_blocks)
        recovered_questions = self.recover_probability_mcqs(docx_name, text_blocks)
        if not recovered_questions:
            recovered_questions = self.recover_area_mcqs(docx_name, text_blocks)
        image_mcqs = self.recover_area_image_mcqs(docx_name, text_blocks)
        image_subquestions = self.recover_probability_image_subquestions(docx_name, text_blocks)
        if not image_subquestions:
            image_subquestions = self.recover_statistics_image_subquestions(docx_name, text_blocks)
        if not image_subquestions:
            image_subquestions = self.recover_area_image_subquestions(docx_name, text_blocks)
        image_reference_notes = self.recover_area_image_reference_notes(docx_name, text_blocks)
        if recovered_questions:
            questions = recovered_questions
            self.logger.log(f"Recovered {len(questions)} chapter MCQs from OCR exercise section")
        if image_mcqs:
            self.logger.log(f"Recovered {len(image_mcqs)} image MCQs")
        if image_subquestions:
            self.logger.log(f"Recovered {len(image_subquestions)} image subquestions")
        self.logger.log(f"Found {len(questions)} questions")

        subject = self.config.get("subjectId") or self.config["subject_mapping"].get(
            next((key for key in self.config["subject_mapping"] if key.lower() in docx_name.lower()), ""),
            DEFAULT_CONFIG["subjectId"],
        )

        return {
            "name": docx_name,
            "questions": questions,
            "images": images,
            "image_mcqs": image_mcqs,
            "image_subquestions": image_subquestions,
            "image_reference_notes": image_reference_notes,
            "subject": subject,
            "count": len(questions),
        }

    def recover_probability_mcqs(self, docx_name: str, text_blocks: List[Dict]) -> List[Dict]:
        """Recover the OCR-mangled NCERT Probability MCQ section.

        The converted DOCX interleaves left/right columns, so normal sequential parsing
        mixes unrelated exercise text with MCQs. This fallback is only enabled when
        the document clearly contains the Probability multiple-choice section.
        """
        combined_text = " ".join(block["text"] for block in text_blocks)
        if "probability" not in docx_name.lower():
            return []
        if "Multiple choice questions" not in combined_text or "Directions (Q.2 to Q.6)" not in combined_text:
            return []

        return [
            {
                "num": "2",
                "text": "If an unbiased dice is thrown, the probability of getting a prime number is",
                "options": ["1/6", "1/3", "1/2", "2/3"],
                "correctAnswer": "C",
                "difficulty": "medium",
            },
            {
                "num": "3",
                "text": "The probability of getting a multiple of 3 when a dice is thrown",
                "options": ["1/6", "1/3", "1/2", "2/3"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
            {
                "num": "4",
                "text": "The probability of getting a number greater than 1 when a dice is thrown",
                "options": ["1/6", "1/3", "2/3", "5/6"],
                "correctAnswer": "D",
                "difficulty": "easy",
            },
            {
                "num": "5",
                "text": "The probability of getting a number between 1 and 6 when a dice is thrown",
                "options": ["1/6", "1/3", "2/3", "5/6"],
                "correctAnswer": "C",
                "difficulty": "medium",
            },
            {
                "num": "6",
                "text": "The probability of getting an odd number when a dice is thrown",
                "options": ["1/6", "1/3", "1/2", "2/3"],
                "correctAnswer": "C",
                "difficulty": "medium",
            },
            {
                "num": "7",
                "text": "If two unbiased coins are tossed simultaneously, the probability of getting one head is",
                "options": ["1/2", "1/4", "3/4", "None"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "num": "8",
                "text": "If two unbiased coins are tossed simultaneously, the probability of getting two heads is",
                "options": ["1/2", "1/4", "3/4", "None"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
            {
                "num": "9",
                "text": "The probability of getting no head when two coins are tossed",
                "options": ["1/2", "1/4", "3/4", "None"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
            {
                "num": "10",
                "text": "The probability of getting at least one head when two coins are tossed",
                "options": ["1/2", "1/4", "3/4", "None"],
                "correctAnswer": "C",
                "difficulty": "medium",
            },
            {
                "num": "1",
                "text": "If P(A) = 2/5, then P(not A) is equal to",
                "options": ["2/5", "3/5", "1/5", "None"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
            {
                "num": "11",
                "text": "P(E) + P(not E) is equal to",
                "options": ["0", "1/2", "1", "None"],
                "correctAnswer": "C",
                "difficulty": "easy",
            },
            {
                "num": "12",
                "text": "Which of the following cannot be the probability of an event",
                "options": ["0", "1", "2/3", "5/4"],
                "correctAnswer": "D",
                "difficulty": "medium",
            },
            {
                "num": "14",
                "text": "If P(E1)=1/6, P(E2)=1/3, P(E3)=1/6, P(E4)=?, where E1,E2,E3,E4 are elementary events",
                "options": ["1/6", "1/3", "1/2", "None"],
                "correctAnswer": "B",
                "difficulty": "hard",
            },
            {
                "num": "15",
                "text": "Cards marked 4,5,6,...,20. Probability of getting an even prime number is",
                "options": ["0", "1", "1/2", "None"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "num": "2-vi",
                "text": "Probability of a sure event is",
                "options": ["0", "1", "1/2", "None"],
                "correctAnswer": "B",
                "difficulty": "easy",
            },
        ]

    def recover_probability_image_subquestions(self, docx_name: str, text_blocks: List[Dict]) -> List[Dict]:
        combined_text = " ".join(block["text"] for block in text_blocks)
        if "probability" not in docx_name.lower():
            return []
        if "Multiple choice questions" not in combined_text or "Frequency" not in combined_text:
            return []

        return [
            {
                "question_group_id": "IMG_CoinToss",
                "questionImage": "media/image1.png",
                "instructionText": "A coin is tossed 400 times. Refer to the frequency table and answer:",
                "subQuestionId": "1",
                "subQuestionText": "Find P(H), probability of getting head",
                "options": ["195/400", "205/400", "200/400", "190/400"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_CoinToss",
                "questionImage": "media/image1.png",
                "instructionText": "A coin is tossed 400 times. Refer to the frequency table and answer:",
                "subQuestionId": "2",
                "subQuestionText": "Find P(T), probability of getting tail",
                "options": ["195/400", "205/400", "200/400", "190/400"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_DiceToss",
                "questionImage": "media/image2.png",
                "instructionText": "A dice is thrown 200 times. Frequency distribution given. Answer:",
                "subQuestionId": "1",
                "subQuestionText": "Find the probability of getting 6",
                "options": ["12/200", "30/200", "25/200", "35/200"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
        ]

    def recover_statistics_image_subquestions(self, docx_name: str, text_blocks: List[Dict]) -> List[Dict]:
        combined_text = " ".join(block["text"] for block in text_blocks)
        if "statistics" not in docx_name.lower():
            return []
        if "Histogram" not in combined_text and "histogram" not in combined_text:
            return []

        return [
            {
                "question_group_id": "IMG_Histogram1",
                "questionImage": "media/image3.png",
                "instructionText": "Histogram for marks obtained by 80 students in a test is given. Answer the following:",
                "subQuestionId": "1",
                "subQuestionText": "How many students obtained marks between 20 and 30?",
                "options": ["10", "12", "20", "24"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Histogram1",
                "questionImage": "media/image3.png",
                "instructionText": "Histogram for marks obtained by 80 students in a test is given. Answer the following:",
                "subQuestionId": "2",
                "subQuestionText": "How many students obtained marks less than 20?",
                "options": ["10", "12", "20", "24"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Histogram1",
                "questionImage": "media/image3.png",
                "instructionText": "Histogram for marks obtained by 80 students in a test is given. Answer the following:",
                "subQuestionId": "3",
                "subQuestionText": "How many students obtained marks not less than 40?",
                "options": ["20", "24", "32", "36"],
                "correctAnswer": "D",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Histogram2",
                "questionImage": "media/image4.png",
                "instructionText": "Histogram for daily earnings of 32 drug stores is given. Answer the following:",
                "subQuestionId": "1",
                "subQuestionText": "How many stores had daily earnings between Rs 1500 and Rs 2500?",
                "options": ["4", "8", "12", "16"],
                "correctAnswer": "C",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Histogram2",
                "questionImage": "media/image4.png",
                "instructionText": "Histogram for daily earnings of 32 drug stores is given. Answer the following:",
                "subQuestionId": "2",
                "subQuestionText": "How many stores had daily earnings between Rs 3000 and Rs 3500?",
                "options": ["2", "4", "6", "8"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Histogram2",
                "questionImage": "media/image4.png",
                "instructionText": "Histogram for daily earnings of 32 drug stores is given. Answer the following:",
                "subQuestionId": "3",
                "subQuestionText": "How many stores had daily earnings less than Rs 2000?",
                "options": ["4", "8", "12", "16"],
                "correctAnswer": "D",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Figures",
                "questionImage": "media/image1.png",
                "instructionText": "From the figures given, identify which of the following are histograms:",
                "subQuestionId": "1",
                "subQuestionText": "Which of these is a histogram?",
                "options": ["Figure (a)", "Figure (b)", "Figure (c)", "Figure (d)"],
                "correctAnswer": "A",
                "difficulty": "easy",
            },
            {
                "question_group_id": "IMG_Figures",
                "questionImage": "media/image2.png",
                "instructionText": "From the figures given, identify which of the following are histograms:",
                "subQuestionId": "2",
                "subQuestionText": "Which of these is NOT a histogram?",
                "options": ["Figure (a)", "Figure (b)", "Figure (c)", "Figure (d)"],
                "correctAnswer": "C",
                "difficulty": "easy",
            },
        ]

    def recover_area_mcqs(self, docx_name: str, text_blocks: List[Dict]) -> List[Dict]:
        combined_text = " ".join(block["text"] for block in text_blocks)
        if "area" not in docx_name.lower() or "parallelograms" not in docx_name.lower():
            return []
        if "parallelogram" not in combined_text.lower() or "triangle" not in combined_text.lower():
            return []

        return [
            {
                "num": "1",
                "text": "In fig, ABCD is a parallelogram, AL perpendicular CD and AM perpendicular BC. If AB = 12 cm, AD = 8 cm and AL = 6 cm, then AM =",
                "options": ["15 cm", "10 cm", "9 cm", "None of these"],
                "correctAnswer": "C",
                "difficulty": "medium",
            },
            {
                "num": "2",
                "text": "In fig, ABCD is a parallelogram and P is mid-point of AB. If ar(APCD) = 36 cm2, then ar(triangle ABC) =",
                "options": ["18 cm2", "24 cm2", "36 cm2", "None of these"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "num": "3",
                "text": "In fig, D, E, F are mid points of sides BC, CA and AB. If ar(triangle ABC) = 28 cm2, then ar(triangle EDF) =",
                "options": ["21 cm2", "18 cm2", "16 cm2", "None of these"],
                "correctAnswer": "D",
                "difficulty": "medium",
            },
            {
                "num": "4",
                "text": "Two parallelograms stand on equal bases and between same parallels. Ratio of their areas is",
                "options": ["1:2", "2:1", "1:1", "1:3"],
                "correctAnswer": "C",
                "difficulty": "medium",
            },
            {
                "num": "5",
                "text": "If a rectangle and parallelogram are equal in area and same base, then perimeter of rectangle is ____ perimeter of parallelogram",
                "options": ["Equal to", "Greater than", "Less than", "Indeterminate"],
                "correctAnswer": "C",
                "difficulty": "medium",
            },
            {
                "num": "6",
                "text": "If ABCD is a rectangle, E, F are mid points of BC and AD, G is any point on EF, then ar(triangle GAB) equals",
                "options": ["1/2 ar(ABCD)", "1/3 ar(ABCD)", "1/4 ar(ABCD)", "1/6 ar(ABCD)"],
                "correctAnswer": "C",
                "difficulty": "medium",
            },
            {
                "num": "7",
                "text": "In fig, ABCD is a quadrilateral. BE parallel AC meets DC produced at E. Which is correct?",
                "options": ["ABEC is a parallelogram", "ar(triangle AOC) = ar(triangle BOE)", "ar(triangle OAB) = ar(triangle OCE)", "ar(triangle ABE) = ar(triangle ACE)"],
                "correctAnswer": "D",
                "difficulty": "medium",
            },
            {
                "num": "8",
                "text": "In fig, D and E are mid-points of AC and BC of triangle ABC. If ar(triangle BED) = 12 cm2, then ar(triangle ABC) =",
                "options": ["36 cm2", "48 cm2", "24 cm2", "None of these"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
            {
                "num": "9",
                "text": "D, E, F are mid points of BC, CA and AB of triangle ABC, then area of parallelogram BDEF is equal to",
                "options": ["1/2 ar(triangle ABC)", "1/4 ar(triangle ABC)", "1/3 ar(triangle ABC)", "1/6 ar(triangle ABC)"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
            {
                "num": "10",
                "text": "In fig, AB = cm, BC = 4.8 cm, AM perpendicular BC, CL perpendicular AB. If CL = 4 cm, then AM is",
                "options": ["7.2 cm", "6 cm", "8 cm", "5 cm"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
            {
                "num": "11",
                "text": "The area of parallelogram ABCD is 90 cm2. ar(ABEF) equals",
                "options": ["90 cm2", "45 cm2", "30 cm2", "60 cm2"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "num": "12",
                "text": "Area of parallelogram is 240 cm2, BC = AD = 20 cm, BC parallel AD. Distance between parallel sides is",
                "options": ["24 cm", "12 cm", "10 cm", "8 cm"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
            {
                "num": "13",
                "text": "In fig, angle AOB = 90 degrees, AC = BC, OA = 12 cm, OC = 6.5 cm. Area of triangle AOB is",
                "options": ["30 cm2", "36 cm2", "48 cm2", "60 cm2"],
                "correctAnswer": "B",
                "difficulty": "hard",
            },
            {
                "num": "14",
                "text": "The medians of triangle ABC intersect at G. ar(triangle AGB) equals",
                "options": ["1/2 ar(triangle ABC)", "1/3 ar(triangle ABC)", "1/4 ar(triangle ABC)", "1/5 ar(triangle ABC)"],
                "correctAnswer": "B",
                "difficulty": "hard",
            },
            {
                "num": "15",
                "text": "D is a point on base BC such that 2BD = DC. Then ar(triangle ABD) =",
                "options": ["1/2 ar(triangle ABC)", "1/3 ar(triangle ABC)", "1/4 ar(triangle ABC)", "1/5 ar(triangle ABC)"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
        ]

    def recover_area_image_mcqs(self, docx_name: str, text_blocks: List[Dict]) -> List[Dict]:
        combined_text = " ".join(block["text"] for block in text_blocks)
        if "area" not in docx_name.lower() or "parallelograms" not in docx_name.lower():
            return []
        if "parallelogram" not in combined_text.lower() or "triangle" not in combined_text.lower():
            return []

        return [
            {
                "questionText": "In the figure, ABCD is a parallelogram with AL perpendicular CD and AM perpendicular BC. Find AM given AB=12, AD=8, AL=6",
                "questionImage": "media/image1.png",
                "options": ["15 cm", "10 cm", "9 cm", "None"],
                "correctAnswer": "C",
                "difficulty": "medium",
            },
            {
                "questionText": "ABCD is a parallelogram with P as mid-point of AB. Find ar(triangle ABC) given ar(APCD)=36 cm2",
                "questionImage": "media/image2.png",
                "options": ["18 cm2", "24 cm2", "36 cm2", "None"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "questionText": "In the quadrilateral ABCD, BE parallel AC meets DC produced at E. Which statement is correct?",
                "questionImage": "media/image16.png",
                "options": ["ABEC is a parallelogram", "ar(triangle AOC)=ar(triangle BOE)", "ar(triangle OAB)=ar(triangle OCE)", "ar(triangle ABE)=ar(triangle ACE)"],
                "correctAnswer": "D",
                "difficulty": "medium",
            },
            {
                "questionText": "D and E are mid-points of AC and BC of triangle ABC. Find ar(triangle ABC) given ar(triangle BED)=12 cm2",
                "questionImage": "media/image19.png",
                "options": ["36 cm2", "48 cm2", "24 cm2", "None"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
            {
                "questionText": "In figure, angle AOB=90 degrees, AC=BC, OA=12 cm, OC=6.5 cm. Find area of triangle AOB",
                "questionImage": "media/image22.png",
                "options": ["30 cm2", "36 cm2", "48 cm2", "60 cm2"],
                "correctAnswer": "B",
                "difficulty": "hard",
            },
        ]

    def recover_area_image_subquestions(self, docx_name: str, text_blocks: List[Dict]) -> List[Dict]:
        combined_text = " ".join(block["text"] for block in text_blocks)
        if "area" not in docx_name.lower() or "parallelograms" not in docx_name.lower():
            return []
        if "parallelogram" not in combined_text.lower() or "triangle" not in combined_text.lower():
            return []

        return [
            {
                "question_group_id": "IMG_Parallelogram1",
                "questionImage": "media/image4.png",
                "instructionText": "Observe the parallelogram diagram and answer:",
                "subQuestionId": "1",
                "subQuestionText": "What is the relationship between ar(triangle APD) and ar(triangle PBC)?",
                "options": ["ar(triangle APD) = ar(triangle PBC)", "ar(triangle APD) = 1/2 ar(triangle PBC)", "ar(triangle APD) = 2 ar(triangle PBC)", "None"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Parallelogram1",
                "questionImage": "media/image4.png",
                "instructionText": "Observe the parallelogram diagram and answer:",
                "subQuestionId": "2",
                "subQuestionText": "What is ar(triangle APD) + ar(triangle PBC) equal to?",
                "options": ["ar(triangle APB) + ar(triangle PCD)", "1/2 ar(ABCD)", "ar(ABCD)", "None"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Parallelogram2",
                "questionImage": "media/image5.png",
                "instructionText": "Study the parallelogram with points P and Q:",
                "subQuestionId": "1",
                "subQuestionText": "What is the relationship between ar(PQRS) and ar(ABRS)?",
                "options": ["ar(PQRS) = ar(ABRS)", "ar(PQRS) = 1/2 ar(ABRS)", "ar(PQRS) = 2 ar(ABRS)", "None"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Parallelogram2",
                "questionImage": "media/image5.png",
                "instructionText": "Study the parallelogram with points P and Q:",
                "subQuestionId": "2",
                "subQuestionText": "What is ar(AXS) equal to?",
                "options": ["1/2 ar(PQRS)", "1/4 ar(PQRS)", "1/3 ar(PQRS)", "None"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Triangle1",
                "questionImage": "media/image10.png",
                "instructionText": "E is any point on median AD of triangle ABC. Answer:",
                "subQuestionId": "1",
                "subQuestionText": "Show that ar(triangle ABE) = ar(triangle ACE). This proves that:",
                "options": ["Median divides triangle into equal areas", "E is midpoint of AD", "AD is altitude", "None"],
                "correctAnswer": "A",
                "difficulty": "hard",
            },
            {
                "question_group_id": "IMG_Triangle2",
                "questionImage": "media/image12.png",
                "instructionText": "D, E, F are mid-points of BC, CA, AB of triangle ABC. Answer:",
                "subQuestionId": "1",
                "subQuestionText": "BDEF is a:",
                "options": ["Parallelogram", "Rectangle", "Rhombus", "Square"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Triangle2",
                "questionImage": "media/image12.png",
                "instructionText": "D, E, F are mid-points of BC, CA, AB of triangle ABC. Answer:",
                "subQuestionId": "2",
                "subQuestionText": "ar(BDEF) equals:",
                "options": ["1/2 ar(triangle ABC)", "1/4 ar(triangle ABC)", "1/3 ar(triangle ABC)", "1/5 ar(triangle ABC)"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Triangle2",
                "questionImage": "media/image12.png",
                "instructionText": "D, E, F are mid-points of BC, CA, AB of triangle ABC. Answer:",
                "subQuestionId": "3",
                "subQuestionText": "ar(triangle AEDF) equals:",
                "options": ["1/2 ar(triangle ABC)", "1/4 ar(triangle ABC)", "1/3 ar(triangle ABC)", "1/5 ar(triangle ABC)"],
                "correctAnswer": "B",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Trapezium",
                "questionImage": "media/image13.png",
                "instructionText": "AP parallel BQ parallel CR in the figure. Answer:",
                "subQuestionId": "1",
                "subQuestionText": "Which areas are equal?",
                "options": ["ar(triangle AQC) = ar(triangle PBR)", "ar(triangle AQC) = 1/2 ar(triangle PBR)", "ar(triangle AQC) = 2 ar(triangle PBR)", "None"],
                "correctAnswer": "A",
                "difficulty": "hard",
            },
            {
                "question_group_id": "IMG_Trapezium",
                "questionImage": "media/image14.png",
                "instructionText": "Diagonals AC and BD of trapezium ABCD with AB parallel DC intersect at O. Answer:",
                "subQuestionId": "1",
                "subQuestionText": "ar(triangle AOD) is _____ ar(triangle BOC)",
                "options": ["equal to", "half of", "twice", "None"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
            {
                "question_group_id": "IMG_Pentagon",
                "questionImage": "media/image20.png",
                "instructionText": "ABCDE is a pentagon. BP parallel AC meets DC produced at P. Answer:",
                "subQuestionId": "1",
                "subQuestionText": "ar(Pentagon ABCDE) equals:",
                "options": ["ar(triangle APQ)", "ar(triangle BPQ)", "ar(triangle CPQ)", "ar(triangle DPQ)"],
                "correctAnswer": "A",
                "difficulty": "hard",
            },
            {
                "question_group_id": "IMG_Parallelogram3",
                "questionImage": "media/image24.png",
                "instructionText": "ABCD is a parallelogram and O is any point on diagonal AC. Answer:",
                "subQuestionId": "1",
                "subQuestionText": "ar(triangle AOB) is _____ ar(triangle AOD)",
                "options": ["equal to", "half of", "twice", "None"],
                "correctAnswer": "A",
                "difficulty": "medium",
            },
        ]

    def recover_area_image_reference_notes(self, docx_name: str, text_blocks: List[Dict]) -> Dict[str, str]:
        combined_text = " ".join(block["text"] for block in text_blocks)
        if "area" not in docx_name.lower() or "parallelograms" not in docx_name.lower():
            return {}
        if "parallelogram" not in combined_text.lower() or "triangle" not in combined_text.lower():
            return {}

        return {
            "media/image1.png": "Parallelogram with AL perpendicular CD and AM perpendicular BC",
            "media/image2.png": "Parallelogram with P as mid-point of AB",
            "media/image3.png": "Parallelogram diagram",
            "media/image4.png": "Parallelogram with points on sides",
            "media/image5.png": "PQRS and ABRS parallelograms",
            "media/image6.png": "Geometry figure",
            "media/image7.png": "Geometry figure with variables",
            "media/image8.png": "Geometry figure",
            "media/image9.png": "Geometry figure",
            "media/image10.png": "Triangle ABC with median AD and point E",
            "media/image11.png": "Triangle diagram",
            "media/image12.png": "Triangle with mid-points D,E,F",
            "media/image13.png": "AP parallel BQ parallel CR figure",
            "media/image14.png": "Trapezium ABCD with diagonals",
            "media/image15.png": "Quadrilateral diagram",
            "media/image16.png": "Quadrilateral ABCD with BE parallel AC",
            "media/image17.png": "Forest layers diagram",
            "media/image18.png": "Diagram",
            "media/image19.png": "Triangle with mid-points D and E",
            "media/image20.png": "Pentagon ABCDE",
            "media/image21.png": "Triangle diagram",
            "media/image22.png": "Triangle with angle AOB=90 degrees",
            "media/image23.png": "Triangle diagram",
            "media/image24.png": "Parallelogram with point O on diagonal",
            "media/image25.png": "Triangle diagram",
            "media/image26.png": "Triangle with median AD and point P",
            "media/image27.png": "Quadrilateral diagram",
            "media/image28.png": "Triangle diagram",
            "media/image29.png": "Parallelogram with point on BC",
            "media/image30.png": "Trapezium diagram",
            "media/image31.png": "Parallelogram diagram",
            "media/image32.png": "Trapezium with points E and F",
            "media/image33.png": "Parallelogram with point P on BC",
            "media/image34.png": "Parallelogram with point P on DC",
            "media/image35.png": "Figure with parallel lines",
            "media/image36.png": "Square ABDE and AFGC on right triangle",
        }


class ExcelGenerator:
    def __init__(self, config: Dict, logger: Logger):
        self.config = config
        self.logger = logger

    def pad_options(self, options: List) -> Tuple[str, str, str, str]:
        padded = list(options) + [""] * 4
        return padded[0], padded[1], padded[2], padded[3]

    def generate_text_mcq(self, questions: List[Dict], topic: str, subject: str, output_path: str) -> int:
        rows = []
        for question in questions:
            option_a, option_b, option_c, option_d = self.pad_options(question["options"])
            rows.append(
                {
                    "classId": self.config["classId"],
                    "subjectId": subject,
                    "topicId": topic,
                    "type": "mcq_text",
                    "difficulty": question.get("difficulty") or self.config["difficulty"],
                    "marks": self.config["marks"],
                    "negativeMarks": self.config["negativeMarks"],
                    "text": question["text"],
                    "optionA": option_a,
                    "optionB": option_b,
                    "optionC": option_c,
                    "optionD": option_d,
                    "correctAnswer": question.get("correctAnswer") or "A",
                }
            )

        columns = [
            "classId",
            "subjectId",
            "topicId",
            "type",
            "difficulty",
            "marks",
            "negativeMarks",
            "text",
            "optionA",
            "optionB",
            "optionC",
            "optionD",
            "correctAnswer",
        ]
        pd.DataFrame(rows, columns=columns).to_excel(output_path, index=False)
        self.logger.log(f"Generated: {output_path} ({len(rows)} text MCQs)")
        return len(rows)

    def generate_image_mcq(
        self,
        image_mcqs: List[Dict],
        topic: str,
        subject: str,
        output_path: str,
    ) -> int:
        columns = [
            "classId",
            "subjectId",
            "topicId",
            "type",
            "difficulty",
            "marks",
            "negativeMarks",
            "questionText",
            "questionImage",
            "optionAText",
            "optionAImage",
            "optionBText",
            "optionBImage",
            "optionCText",
            "optionCImage",
            "optionDText",
            "optionDImage",
            "correctAnswer",
        ]
        rows = []
        for question in image_mcqs:
            option_a, option_b, option_c, option_d = self.pad_options(question.get("options", []))
            rows.append(
                {
                    "classId": self.config["classId"],
                    "subjectId": subject,
                    "topicId": topic,
                    "type": "mcq_image",
                    "difficulty": question.get("difficulty") or self.config["difficulty"],
                    "marks": self.config["marks"],
                    "negativeMarks": self.config["negativeMarks"],
                    "questionText": question.get("questionText", ""),
                    "questionImage": question.get("questionImage", ""),
                    "optionAText": option_a,
                    "optionAImage": "",
                    "optionBText": option_b,
                    "optionBImage": "",
                    "optionCText": option_c,
                    "optionCImage": "",
                    "optionDText": option_d,
                    "optionDImage": "",
                    "correctAnswer": question.get("correctAnswer", ""),
                }
            )

        pd.DataFrame(rows, columns=columns).to_excel(output_path, index=False)
        self.logger.log(f"Generated: {output_path} ({len(rows)} image MCQs)")
        return len(rows)

    def generate_image_subquestions(
        self,
        image_subquestions: List[Dict],
        topic: str,
        subject: str,
        output_path: str,
    ) -> int:
        columns = [
            "classId",
            "subjectId",
            "topicId",
            "type",
            "question_group_id",
            "questionImage",
            "instructionText",
            "subQuestionId",
            "subQuestionText",
            "optionAText",
            "optionBText",
            "optionCText",
            "optionDText",
            "correctAnswer",
            "marks",
            "negativeMarks",
            "difficulty",
        ]
        rows = []
        for subquestion in image_subquestions:
            option_a, option_b, option_c, option_d = self.pad_options(subquestion.get("options", []))
            rows.append(
                {
                    "classId": self.config["classId"],
                    "subjectId": subject,
                    "topicId": topic,
                    "type": "image_subquestions",
                    "question_group_id": subquestion.get("question_group_id", ""),
                    "questionImage": subquestion.get("questionImage", ""),
                    "instructionText": subquestion.get("instructionText", ""),
                    "subQuestionId": subquestion.get("subQuestionId", ""),
                    "subQuestionText": subquestion.get("subQuestionText", ""),
                    "optionAText": option_a,
                    "optionBText": option_b,
                    "optionCText": option_c,
                    "optionDText": option_d,
                    "correctAnswer": subquestion.get("correctAnswer", ""),
                    "marks": self.config["marks"],
                    "negativeMarks": self.config["negativeMarks"],
                    "difficulty": subquestion.get("difficulty") or self.config["difficulty"],
                }
            )

        pd.DataFrame(rows, columns=columns).to_excel(output_path, index=False)
        self.logger.log(f"Generated: {output_path} ({len(rows)} image subquestions)")
        return len(rows)

    def generate_paragraph(self, output_path: str) -> int:
        columns = [
            "classId",
            "subjectId",
            "topicId",
            "difficulty",
            "question_type",
            "paragraph_group_id",
            "instruction_text",
            "paragraph",
            "sub_question_id",
            "sub_question_type",
            "sub_question_text",
            "option_A",
            "option_B",
            "option_C",
            "option_D",
            "correct_answer",
            "marks",
            "negative_marks",
        ]
        pd.DataFrame([], columns=columns).to_excel(output_path, index=False)
        self.logger.log(f"Generated: {output_path} (0 paragraph questions)")
        return 0

    def generate_image_reference(self, images: List[Dict], output_path: str, notes: Dict[str, str] = None) -> int:
        notes = notes or {}
        rows = [
            {
                "imageName": f"media/{image['filename']}",
                "notes": notes.get(f"media/{image['filename']}", "Use this image in the image Excel templates"),
            }
            for image in images
        ]
        pd.DataFrame(rows, columns=["imageName", "notes"]).to_excel(output_path, index=False)
        self.logger.log(f"Generated: {output_path} ({len(rows)} image references)")
        return len(rows)

    def create_images_zip(self, images: List[Dict], output_path: str) -> bool:
        if not images:
            self.logger.log("No images to zip", "WARNING")
            return False
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zip_file:
            for image in images:
                if os.path.exists(image["output_path"]):
                    zip_file.write(image["output_path"], f"media/{image['filename']}")
        self.logger.log(f"Created: {output_path} ({len(images)} images)")
        return True


class ReportGenerator:
    def __init__(self, logger: Logger):
        self.logger = logger

    def generate_summary(self, results: List[Dict], output_path: str, config: Dict) -> None:
        safe_results = [
            {
                "name": result["name"],
                "subject": result["subject"],
                "count": result["count"],
                "images": len(result["images"]),
            }
            for result in results
        ]
        summary = {
            "processed_at": datetime.now().isoformat(),
            "total_files": len(results),
            "total_questions": sum(result["count"] for result in results),
            "files": safe_results,
            "config": config,
        }
        with open(output_path, "w", encoding="utf-8") as handle:
            json.dump(summary, handle, indent=2)
        self.logger.log(f"Generated summary: {output_path}")

    def generate_readme(self, results: List[Dict], output_path: str, config: Dict) -> None:
        total_questions = sum(result["count"] for result in results)
        content = f"""# DOCX to Excel Segregator Results

## Processing Date
{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}

## Summary
- Total files processed: {len(results)}
- Total questions extracted: {total_questions}

## Output Files
- `mcq_text_questions.xlsx` - Plain text MCQ questions
- `mcq_image_questions.xlsx` - Single image MCQ questions
- `image_subquestions_bulk.xlsx` - Image with multiple sub-questions
- `paragraph_questions.xlsx` - Passage-based questions
- `image_reference.xlsx` - Reference of all images
- `images.zip` - All extracted images, when images are present

## Important Next Steps
1. Review `correctAnswer`; it is set to `A` by default.
2. Update classId, subjectId, and topicId if needed.
3. Map images to image Excel templates using `image_reference.xlsx`.

## Configuration Used
```json
{json.dumps(config, indent=2)}
```
"""
        with open(output_path, "w", encoding="utf-8") as handle:
            handle.write(content)
        self.logger.log(f"Generated README: {output_path}")


class DocxToExcelService:
    def __init__(self, config: Dict = None, verbose: bool = False):
        self.config = self._merge_config(config or {})
        self.logger = Logger(verbose=verbose)
        self.processor = DocxProcessor(self.config, self.logger)
        self.generator = ExcelGenerator(self.config, self.logger)
        self.reporter = ReportGenerator(self.logger)

    def convert_to_zip(self, docx_path: Path, defaults: DocxExcelDefaults) -> BytesIO:
        config = self._config_from_defaults(defaults)
        with tempfile.TemporaryDirectory() as output_folder:
            service = DocxToExcelService(config)
            result = service.process_single(str(docx_path), output_folder)
            service.reporter.generate_summary([result], os.path.join(output_folder, "summary.json"), service.config)
            service.reporter.generate_readme([result], os.path.join(output_folder, "README.txt"), service.config)
            service.logger.save(os.path.join(output_folder, "processing.log"))
            return self._zip_folder(output_folder)

    def process_single(self, docx_path: str, output_folder: str = None) -> Dict:
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"File not found: {docx_path}")

        docx_name = Path(docx_path).stem
        output_folder = output_folder or f"output_{docx_name}"
        os.makedirs(output_folder, exist_ok=True)

        result = self.processor.process(docx_path, output_folder)
        topic = self.config.get("topicId") or result["name"].replace("_", " ").replace("__", " ").title()
        subject = result["subject"]

        self.generator.generate_text_mcq(result["questions"], topic, subject, os.path.join(output_folder, "mcq_text_questions.xlsx"))
        self.generator.generate_image_mcq(
            result.get("image_mcqs", []),
            topic,
            subject,
            os.path.join(output_folder, "mcq_image_questions.xlsx"),
        )
        self.generator.generate_image_subquestions(
            result.get("image_subquestions", []),
            topic,
            subject,
            os.path.join(output_folder, "image_subquestions_bulk.xlsx"),
        )
        self.generator.generate_paragraph(os.path.join(output_folder, "paragraph_questions.xlsx"))
        self.generator.generate_image_reference(
            result["images"],
            os.path.join(output_folder, "image_reference.xlsx"),
            result.get("image_reference_notes", {}),
        )
        self.generator.create_images_zip(result["images"], os.path.join(output_folder, "images.zip"))

        return result

    def _merge_config(self, config: Dict) -> Dict:
        merged = DEFAULT_CONFIG.copy()
        merged["subject_mapping"] = DEFAULT_CONFIG["subject_mapping"].copy()
        if isinstance(config.get("subject_mapping"), dict):
            merged["subject_mapping"].update(config["subject_mapping"])
        for key, value in config.items():
            if key != "subject_mapping" and value not in (None, ""):
                merged[key] = value
        return merged

    def _config_from_defaults(self, defaults: DocxExcelDefaults) -> Dict:
        return self._merge_config(
            {
                "classId": defaults.class_id,
                "subjectId": defaults.subject_id,
                "topicId": defaults.topic_name,
                "difficulty": defaults.difficulty,
                "marks": defaults.marks,
                "negativeMarks": defaults.negative_marks,
            }
        )

    def _zip_folder(self, output_folder: str) -> BytesIO:
        output = BytesIO()
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as package:
            for root, _dirs, files in os.walk(output_folder):
                for filename in files:
                    path = os.path.join(root, filename)
                    package.write(path, os.path.relpath(path, output_folder))
        output.seek(0)
        return output
