from __future__ import annotations

import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Optional, Union

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt


@dataclass(frozen=True)
class ConversionSettings:
    top_margin_cm: float = 1.5
    bottom_margin_cm: float = 1.5
    left_margin_cm: float = 2
    right_margin_cm: float = 2
    default_font_name: str = "Calibri"
    default_font_size: int = 11
    min_image_width: int = 35
    min_image_height: int = 35
    render_scale: float = 2.5


class PdfToDocxService:
    def __init__(self, settings: Optional[ConversionSettings] = None) -> None:
        self.settings = settings or ConversionSettings()

    def convert(
        self,
        pdf_path: Union[str, Path],
        output_docx_path: Union[str, Path],
        image_dir: Optional[Union[str, Path]] = None,
    ) -> Path:
        pdf_path = Path(pdf_path)
        output_docx_path = Path(output_docx_path)

        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")

        output_docx_path.parent.mkdir(parents=True, exist_ok=True)

        if image_dir is not None:
            permanent_image_dir = Path(image_dir)
            permanent_image_dir.mkdir(parents=True, exist_ok=True)
            return self._convert_with_image_dir(pdf_path, output_docx_path, permanent_image_dir)

        with tempfile.TemporaryDirectory() as temp_image_dir:
            return self._convert_with_image_dir(pdf_path, output_docx_path, Path(temp_image_dir))

    def _convert_with_image_dir(
        self,
        pdf_path: Path,
        output_docx_path: Path,
        image_dir: Path,
    ) -> Path:
        word_doc = Document()
        self._configure_document(word_doc)

        pdf_doc = fitz.open(str(pdf_path))
        try:
            self._write_pdf_content(pdf_doc, word_doc, image_dir)
        finally:
            pdf_doc.close()

        word_doc.save(str(output_docx_path))

        return output_docx_path

    def _configure_document(self, word_doc: Document) -> None:
        for section in word_doc.sections:
            section.top_margin = Cm(self.settings.top_margin_cm)
            section.bottom_margin = Cm(self.settings.bottom_margin_cm)
            section.left_margin = Cm(self.settings.left_margin_cm)
            section.right_margin = Cm(self.settings.right_margin_cm)

        style = word_doc.styles["Normal"]
        style.font.name = self.settings.default_font_name
        style.font.size = Pt(self.settings.default_font_size)

    def _write_pdf_content(
        self,
        pdf_doc: fitz.Document,
        word_doc: Document,
        image_dir: Path,
    ) -> None:
        image_counter = 0

        for page_index in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_index)
            blocks = page.get_text("dict")["blocks"]
            blocks = sorted(blocks, key=lambda block: block["bbox"][1])

            for block in blocks:
                if block["type"] == 0:
                    self._add_text_block(word_doc, block)
                elif block["type"] == 1:
                    image_counter = self._add_image_block(
                        word_doc=word_doc,
                        page=page,
                        block=block,
                        image_dir=image_dir,
                        page_index=page_index,
                        image_counter=image_counter,
                    )

            if page_index < len(pdf_doc) - 1:
                word_doc.add_page_break()

    def _add_text_block(self, word_doc: Document, block: dict) -> None:
        full_text = ""
        font_size = self.settings.default_font_size
        is_bold = False

        for line in block.get("lines", []):
            for span in line.get("spans", []):
                full_text += span["text"] + " "
                if span.get("size"):
                    font_size = span["size"]
                if span.get("flags", 0) & 16:
                    is_bold = True

        full_text = full_text.strip()
        if not full_text:
            return

        paragraph = word_doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(3)

        run = paragraph.add_run(full_text)
        run.font.size = Pt(max(10, min(round(font_size), 14)))
        run.bold = is_bold

    def _add_image_block(
        self,
        word_doc: Document,
        page: fitz.Page,
        block: dict,
        image_dir: Path,
        page_index: int,
        image_counter: int,
    ) -> int:
        x0, y0, x1, y1 = block["bbox"]
        width = x1 - x0
        height = y1 - y0

        if width < self.settings.min_image_width or height < self.settings.min_image_height:
            return image_counter

        image_path = image_dir / f"page_{page_index + 1}_img_{image_counter}.png"
        clip = fitz.Rect(block["bbox"])
        pix = page.get_pixmap(
            matrix=fitz.Matrix(self.settings.render_scale, self.settings.render_scale),
            clip=clip,
        )
        pix.save(str(image_path))

        page_width = page.rect.width
        relative_width = width / page_width
        display_width = min(5.8, max(2.0, relative_width * 6.5))

        paragraph = word_doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)

        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(display_width))

        return image_counter + 1
